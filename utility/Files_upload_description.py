import os
import json
import re
import logging
from itertools import chain
from io import BytesIO
from dotenv import load_dotenv
from bson import ObjectId
from openai import OpenAI
from google.cloud import storage
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from langchain.text_splitter import CharacterTextSplitter

# Import the centralized mongo_crud function
from Databases.mongo import mongo_crud

# Setup logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# Load environment variables from .env file
load_dotenv()

# Retrieve OpenAI API Key from environment variables
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    raise ValueError("OpenAI API key is not set. Please set it in the .env file.")

# Initialize OpenAI client
client = OpenAI(api_key=openai_api_key)

# Initialize the text splitter
text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

# Database and collection constants
#DB_NAME = "ChatbotDB-DEV"
COLLECTION_FILES = "files"

def mongo_operation(operation, collection_name=COLLECTION_FILES, query=None, update=None, start=0, stop=10):
    """Wrapper to call mongo_crud with fixed DB name, no need for host/port."""
    return mongo_crud(
       
        collection_name=collection_name,
        operation=operation,
        query=query or {},
        update=update or {},
        start=start,
        stop=stop
    )

def generate_openai_output(text):
    prompt = f"""Please read the following text or pdf, excel, doc, etc data and write a description, keywords and tags as given below provide the response in JSON format:
{{
    "description": "10-15 word description of the content",
    "keywords": ["5 important keywords"],
    "tags": ["relevant tags such as product name, purpose"]
}}

{text}"""

    response = client.chat.completions.create(
       model="gpt-4o-mini",

        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    content = response.choices[0].message.content.strip()
    # print(content)
    return content

def read_pdf_from_gcs(bucket_name, blob_names, chatbot_id, version_id):
    """Read PDFs from GCS and extract text with error handling"""
    complete_document = []
    print(f"Processing the following blobs: {blob_names}")

    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)

    for blob_name in blob_names:
        try:
            logger.info(f"Processing blob: {blob_name}")
            blob = bucket.blob(blob_name)
            pdf_bytes = blob.download_as_bytes()
            pdf_file = BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            data = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    data.append(text)
                else:
                    logger.warning(f"No text extracted from page {page_num} in blob {blob_name}")

            pdf_text = '. '.join(data)

            file_des = generate_openai_output(pdf_text)
            file_des = json.loads(file_des)

            description = file_des.get('description', 'No description')
            keywords = file_des.get('keywords', [])
            tags = file_des.get('tags', [])

            print(f"Generated description for {blob_name}: {description}")
            print(f"Generated keywords for {blob_name}: {keywords}")
            print(f"Generated tags for {blob_name}: {tags}")

            logger.info(f"Generated description: {description}")
            logger.info(f"Generated keywords: {keywords}")
            logger.info(f"Generated tags: {tags}")

            chatbot_oid = ObjectId(chatbot_id)
            version_oid = ObjectId(version_id)

            mongo_operation(
                operation="update",
                collection_name=COLLECTION_FILES,
                query={"chatbot_id": chatbot_oid, "version_id": version_oid, "file_name": blob_name},
                update={"description": description, "keywords": keywords, "tags": tags}
            )
            logger.info(f"MongoDB updated for blob: {blob_name}")

        except Exception as e:
            logger.error(f"Error processing blob {blob_name}: {e}")

    flattened_docs = list(chain.from_iterable(complete_document))
    return flattened_docs

def document_splitter(text: str):
    try:
        from langchain_core.documents import Document
        long_doc = [Document(page_content=text)]
        docs = text_splitter.split_documents(long_doc)
        return docs
    except Exception as e:
        logger.error(f"Document splitter failed: {str(e)}")
        raise

def read_documents_from_gcs(bucket_name, blob_names, chatbot_id, version_id):
    """Read documents from GCS, extract text, enrich metadata with OpenAI, and update MongoDB"""
    try:
        all_docs = []
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)

        for blob_name in blob_names:
            logger.info(f"📂 Processing blob: {blob_name}")
            blob = bucket.blob(blob_name)

            if not blob.exists():
                logger.warning(f"⚠️ Blob '{blob_name}' not found in bucket '{bucket_name}'. Skipping.")
                continue

            try:
                file_bytes = blob.download_as_bytes()
                ext = os.path.splitext(blob_name)[-1].lower()
                logger.info(f"📄 Detected file type: {ext}")

                if ext == '.pdf':
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                    text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])

                elif ext == '.txt':
                    text = file_bytes.decode('utf-8', errors='ignore')

                elif ext == '.docx':
                    doc = DocxDocument(BytesIO(file_bytes))
                    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

                elif ext in ['.xls', '.xlsx']:
                    excel = pd.read_excel(BytesIO(file_bytes), sheet_name=None)
                    text = "\n\n".join([f"Sheet: {sheet}\n{df.to_string(index=False)}" for sheet, df in excel.items()])

                else:
                    logger.warning(f"❌ Unsupported file format: {ext} for '{blob_name}'. Skipping.")
                    continue

                docs = document_splitter(text)
                all_docs.append(docs)

                openai_response = generate_openai_output(docs)
                try:
                    cleaned_category = re.sub(r'```json|```', '', openai_response).strip()
                    file_meta = json.loads(cleaned_category)
                    description = file_meta.get('description', 'No description')
                    keywords = file_meta.get('keywords', [])
                    tags = file_meta.get('tags', [])
                except Exception as json_err:
                    logger.warning(f"⚠️ Failed to parse OpenAI response: {openai_response}")
                    description, keywords, tags = 'No description', [], []

                logger.info(f"📝 Description for {blob_name}: {description}")
                logger.info(f"🔑 Keywords: {keywords}")
                logger.info(f"🏷️ Tags: {tags}")

                mongo_operation(
                    operation="update",
                    collection_name=COLLECTION_FILES,
                    query={
                        "chatbot_id": ObjectId(chatbot_id),
                        "version_id": ObjectId(version_id),
                        "file_name": blob_name
                    },
                    update={"description": description, "keywords": keywords, "tags": tags}
                )
                logger.info(f"✅ MongoDB updated for blob: {blob_name}")

            except Exception as file_err:
                logger.error(f"❌ Error processing blob {blob_name}: {str(file_err)}")
                continue

        return list(chain.from_iterable(all_docs))

    except Exception as e:
        logger.error(f"🔥 Document reading pipeline failed: {str(e)}")
        raise

def description_from_gcs(bucket_name, blob_names, chatbot_id, version_id):
    """Load various documents from GCS and return extracted documents"""
    try:
        docs = read_documents_from_gcs(bucket_name, blob_names, chatbot_id, version_id)
        if not docs:
            logger.warning("No documents were extracted.")
            return "No documents extracted."
        return docs
    except Exception as e:
        logger.error(f"Failed to process GCS files: {e}")
        return f"Error processing GCS files: {e}"
